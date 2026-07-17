import streamlit as st
import cv2
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont, ImageEnhance
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile

# --- PAGE SETUP ---
st.set_page_config(
    page_title="AI Watermark Studio",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("🛡️ AI Watermark Studio")
st.markdown("Add text and logo watermarks to Images, Videos, PDFs, and Word Documents.")

# --- UTILITY FUNCTIONS ---
def get_watermark_position(img_w, img_h, wm_w, wm_h, position, custom_x=0, custom_y=0):
    padding = 20
    if position == "Top Left":
        return (padding, padding)
    elif position == "Top Right":
        return (img_w - wm_w - padding, padding)
    elif position == "Bottom Left":
        return (padding, img_h - wm_h - padding)
    elif position == "Bottom Right":
        return (img_w - wm_w - padding, img_h - wm_h - padding)
    elif position == "Center":
        return ((img_w - wm_w) // 2, ((img_h - wm_h) // 2))
    else:  # Custom
        return (int(custom_x), int(custom_y))

def apply_image_opacity(im, opacity):
    """Applies opacity to a PIL image with an alpha channel."""
    alpha = im.split()[3]
    alpha = ImageEnhance.Brightness(alpha).enhance(opacity)
    im.putalpha(alpha)
    return im

# --- CORE WATERMARKING ENGINES ---

def watermark_image(src_image, wm_type, text_str, logo_file, opacity, font_size, color_hex, rotation, position, tile, cx, cy):
    # Convert incoming PIL image to RGBA
    base = src_image.convert("RGBA")
    txt_layer = Image.new("RGBA", base.size, (255, 255, 255, 0))
    
    # Create the watermark item
    if wm_type == "Text":
        # Safe fallback font handling
        try:
            font = ImageFont.truetype("arial.ttf", font_size)
        except:
            font = ImageFont.load_default()
            
        # Parse hex color
        h = color_hex.lstrip('#')
        rgb = tuple(int(h[i:i+2], 16) for i in (0, 2, 4))
        fill_color = rgb + (int(opacity * 255),)
        
        # Draw on an isolated canvas to allow precise rotation scaling
        dummy = Image.new("RGBA", (base.width, base.height))
        draw = ImageDraw.Draw(dummy)
        
        # Get text size
        bbox = draw.textbbox((0, 0), text_str, font=font)
        wm_w, wm_h = bbox[2] - bbox[0], bbox[3] - bbox[1]
        
        # Create tight-fitting canvas for the text block to rotate cleanly
        text_canvas = Image.new("RGBA", (wm_w + 20, wm_h + 20), (0,0,0,0))
        text_draw = ImageDraw.Draw(text_canvas)
        text_draw.text((10, 10), text_str, font=font, fill=fill_color)
        text_canvas = text_canvas.rotate(rotation, expand=True)
        wm_w, wm_h = text_canvas.size
    else:
        if logo_file is None:
            return src_image
        logo = Image.open(logo_file).convert("RGBA")
        # Resize logo reasonably based on image size
        scale_factor = font_size / 100.0  # Use font size slider as scaling gauge
        wm_w = int(base.width * 0.2 * scale_factor)
        wm_h = int(logo.height * (wm_w / logo.width))
        text_canvas = logo.resize((wm_w, wm_h))
        text_canvas = apply_image_opacity(text_canvas, opacity)
        text_canvas = text_canvas.rotate(rotation, expand=True)
        wm_w, wm_h = text_canvas.size

    if tile:
        for x in range(0, base.width, wm_w + 100):
            for y in range(0, base.height, wm_h + 100):
                txt_layer.paste(text_canvas, (x, y), text_canvas)
    else:
        pos = get_watermark_position(base.width, base.height, wm_w, wm_h, position, cx, cy)
        txt_layer.paste(text_canvas, pos, text_canvas)
        
    return Image.alpha_composite(base, txt_layer).convert("RGB")

def watermark_video(video_path, out_path, wm_type, text_str, logo_file, opacity, font_size, color_hex, position, cx, cy):
    cap = cv2.VideoCapture(video_path)
    width  = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    fps    = cap.get(cv2.CAP_PROP_FPS)
    fourcc = cv2.VideoWriter_fourcc(*'mp4v')
    out = cv2.VideoWriter(out_path, fourcc, fps, (width, height))
    
    logo_img = None
    if wm_type == "Logo" and logo_file is not None:
        logo_img = Image.open(logo_file).convert("RGBA")
        scale_factor = font_size / 100.0
        wm_w = int(width * 0.2 * scale_factor)
        wm_h = int(logo_img.height * (wm_w / logo_img.width))
        logo_img = logo_img.resize((wm_w, wm_h))
        logo_img = apply_image_opacity(logo_img, opacity)

    # Performance limit processing to first 150 frames for interactive responsiveness
    frame_count = 0
    while cap.isOpened() and frame_count < 150:
        ret, frame = cap.get()
        if not ret:
            break
        
        # Convert CV2 frame (BGR) to PIL (RGB)
        frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        pil_img = Image.fromarray(frame_rgb)
        
        # Re-use our robust Image Watermarking Pipeline
        processed_pil = watermark_image(
            pil_img, wm_type, text_str, logo_file, opacity, 
            font_size, color_hex, 0, position, False, cx, cy
        )
        
        # Back to Open CV Matrix format
        frame_out = cv2.cvtColor(np.array(processed_pil), cv2.COLOR_RGB2BGR)
        out.write(frame_out)
        frame_count += 1
        
    cap.release()
    out.release()

def watermark_pdf(pdf_path, out_path, wm_type, text_str, logo_file, opacity, font_size, color_hex, diagonal):
    doc = fitz.open(pdf_path)
    
    # Save a temporary visual asset if using image-based logo layout
    temp_img_path = None
    if wm_type == "Logo" and logo_file is not None:
        logo = Image.open(logo_file).convert("RGBA")
        logo = apply_image_opacity(logo, opacity)
        tf = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        logo.save(tf.name)
        temp_img_path = tf.name

    h = color_hex.lstrip('#')
    rgb = tuple(int(h[i:i+2], 16) / 255.0 for i in (0, 2, 4))

    for page in doc:
        rect = page.rect
        if wm_type == "Text":
            # Direct matrix rotation setup for targeted diagonal placement
            p = fitz.Point(rect.width * 0.1, rect.height * 0.5) if diagonal else fitz.Point(rect.width * 0.3, rect.height * 0.8)
            matrix = fitz.Matrix(45 if diagonal else 0)
            page.insert_text(
                p, text_str, fontsize=font_size, 
                color=rgb, fill_opacity=opacity, rotate=45 if diagonal else 0
            )
        else:
            if temp_img_path:
                # Target dead center layout scaling footprint area
                w_box = rect.width * 0.4
                h_box = rect.height * 0.3
                target_rect = fitz.Rect(
                    (rect.width - w_box)/2, (rect.height - h_box)/2,
                    (rect.width + w_box)/2, (rect.height + h_box)/2
                )
                page.insert_image(target_rect, filename=temp_img_path)

    doc.save(out_path)
    doc.close()
    if temp_img_path and os.path.exists(temp_img_path):
        os.remove(temp_img_path)

def watermark_docx(docx_path, out_path, wm_type, text_str, logo_file):
    doc = Document(docx_path)
    
    if wm_type == "Text":
        for section in doc.sections:
            header = section.header
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f" [{text_str}] ")
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.bold = True
    else:
        if logo_file is not None:
            # Document header stream structure layout manipulation
            tf = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            Image.open(logo_file).save(tf.name)
            for section in doc.sections:
                header = section.header
                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.add_run().add_picture(tf.name, width=Inches(1.0))
            os.remove(tf.name)

    doc.save(out_path)

# --- SIDEBAR INTERFACE ARCHITECTURE ---
st.sidebar.header("🎨 Watermark Customization")
wm_type = st.sidebar.radio("Watermark Variant", ["Text", "Logo"])

text_str = ""
logo_file = None

if wm_type == "Text":
    text_str = st.sidebar.text_input("Watermark Wording String", "CONFIDENTIAL")
    color_hex = st.sidebar.color_picker("Font Accent Hex", "#FF0000")
else:
    logo_file = st.sidebar.file_uploader("Upload Transparent Brand Graphic Asset", type=["png", "jpg", "jpeg", "webp"])
    color_hex = "#000000" # Fallback

opacity = st.sidebar.slider("Opacity Density", 0.1, 1.0, 0.4, step=0.1)
font_size = st.sidebar.slider("Scale Sizing Weight", 10, 200, 50, step=5)
rotation = st.sidebar.slider("Rotation Orientation (Degrees)", -180, 180, 0, step=5)

st.sidebar.subheader("Positioning Parameters")
tile = st.sidebar.checkbox("Tile Mode Pattern Layout", value=False)

position = "Center"
cx, cy = 0, 0
if not tile:
    position = st.sidebar.selectbox("Anchor Grid Position", ["Center", "Top Left", "Top Right", "Bottom Left", "Bottom Right", "Custom Coordinates"])
    if position == "Custom Coordinates":
        cx = st.sidebar.number_input("Absolute X Offset Pixel Position", min_value=0, value=100)
        cy = st.sidebar.number_input("Absolute Y Offset Pixel Position", min_value=0, value=100)

# --- DOCUMENT & VIEWPORT APPLICATION CONTENT TABS ---
tab_img, tab_vid, tab_pdf, tab_docx = st.tabs(["🖼️ Images", "🎥 Video Assets", "📄 PDF Documents", "📝 Word Documents"])

# 1. IMAGE INTERFACE PROCESSING PIPELINE
with tab_img:
    st.header("Image Pipeline Workspace")
    img_files = st.file_uploader("Drop target images layout workspace", type=["png", "jpg", "jpeg", "webp"], accept_multiple_files=True)
    
    if img_files:
        for idx, img_file in enumerate(img_files):
            orig_img = Image.open(img_file)
            processed_img = watermark_image(
                orig_img, wm_type, text_str, logo_file, opacity,
                font_size, color_hex, rotation, position, tile, cx, cy
            )
            
            # Direct UI Matrix preview display matching column space allocations
            col1, col2 = st.columns(2)
            with col1:
                st.image(orig_img, caption=f"Original Profile: {img_file.name}", use_container_width=True)
            with col2:
                st.image(processed_img, caption=f"Watermarked Target: {img_file.name}", use_container_width=True)
                
            # Temporary disk allocation sequence buffer preservation
            tmp_img_out = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            processed_img.save(tmp_img_out.name, format="PNG")
            with open(tmp_img_out.name, "rb") as f:
                st.download_button(
                    label=f"Download Processed Image: {img_file.name}",
                    data=f.read(),
                    file_name=f"watermarked_{img_file.name}",
                    mime="image/png",
                    key=f"btn_img_{idx}"
                )
            os.remove(tmp_img_out.name)

# 2. VIDEO INTERFACE PROCESSING PIPELINE
with tab_vid:
    st.header("Video Engine Asset Workspace")
    st.info("💡 Notice: To maximize cloud execution efficiency, web-rendering steps isolate processing validation to the first 150 frames.")
    vid_file = st.file_uploader("Upload Target Video Clip Asset", type=["mp4", "avi", "mov", "mkv", "webm"])
    
    if vid_file:
        t_in = tempfile.NamedTemporaryFile(delete=False, suffix=f".{vid_file.name.split('.')[-1]}")
        t_in.write(vid_file.read())
        t_in.close()
        
        t_out = tempfile.NamedTemporaryFile(delete=False, suffix=".mp4")
        t_out.close()
        
        with st.spinner("Executing CV2 frame translation sequences..."):
            watermark_video(t_in.name, t_out.name, wm_type, text_str, logo_file, opacity, font_size, color_hex, position, cx, cy)
            
        st.success("Rendering run complete.")
        with open(t_out.name, "rb") as f:
            st.download_button(
                label="Download Processed Video Output",
                data=f.read(),
                file_name=f"watermarked_{vid_file.name}",
                mime="video/mp4"
            )
            
        os.remove(t_in.name)
        os.remove(t_out.name)

# 3. PDF INTERFACE PROCESSING PIPELINE
with tab_pdf:
    st.header("PDF Layout Engineering Workspace")
    diagonal_pdf = st.checkbox("Rotate Diagonal Overlay Placement Direction Matrix", value=True)
    pdf_file = st.file_uploader("Upload Target Document PDF Schema", type=["pdf"])
    
    if pdf_file:
        t_pdf_in = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        t_pdf_in.write(pdf_file.read())
        t_pdf_in.close()
        
        t_pdf_out = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        t_pdf_out.close()
        
        with st.spinner("Processing PyMuPDF layer canvas vectors..."):
            watermark_pdf(t_pdf_in.name, t_pdf_out.name, wm_type, text_str, logo_file, opacity, font_size, color_hex, diagonal_pdf)
            
        st.success("Target PDF Watermarking Sequence Executed Successfully.")
        with open(t_pdf_out.name, "rb") as f:
            st.download_button(
                label="Download Finished PDF Asset Architecture",
                data=f.read(),
                file_name=f"watermarked_{pdf_file.name}",
                mime="application/pdf"
            )
        os.remove(t_pdf_in.name)
        os.remove(t_pdf_out.name)

# 4. WORD INTERFACE PROCESSING PIPELINE
with tab_docx:
    st.header("Microsoft Word Header Processing Module")
    docx_file = st.file_uploader("Upload Document File Block (.docx format)", type=["docx"])
    
    if docx_file:
        t_doc_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        t_doc_in.write(docx_file.read())
        t_doc_in.close()
        
        t_doc_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        t_doc_out.close()
        
        with st.spinner("Injecting XML Header Stream Components..."):
            watermark_docx(t_doc_in.name, t_doc_out.name, wm_type, text_str, logo_file)
            
        st.success("Word Document Assembly Processing Terminated Nicely.")
        with open(t_doc_out.name, "rb") as f:
            st.download_button(
                label="Download Completed Word Document Structure",
                data=f.read(),
                file_name=f"watermarked_{docx_file.name}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        os.remove(t_doc_in.name)
        os.remove(t_doc_out.name)
